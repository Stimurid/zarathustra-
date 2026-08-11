"""B-5.5 Веха 5 — File attach extractor + normalizer.

Пользователь через UI прикрепляет файл (MD/TXT/JSON) во время идущего рана
через WS/REST intervention kind=attach_file. Payload доходит до
Pipeline._consume_pending, оттуда — сюда для нормализации, а Pipeline
инжектит normalized text в system-prompt следующего persona turn'а
(или в user context, если attach_to_persona не задан).

Ограничения (MVP):
  - Text-only форматы: .md, .txt, .json, .yaml, .yml.
  - Максимум ATTACH_MAX_CHARS (default 30 000) — обрезается с пометкой.
  - Одна attach = ≤5 KB инжектится (первый chunk); полный текст в audit.
  - PDF/DOCX/HTML — не поддерживается в MVP (сообщение об ошибке).

Нет побочных эффектов: чистая функция normalize().
"""
from __future__ import annotations

import json as _json
from dataclasses import dataclass
from typing import Any


ATTACH_MAX_CHARS = 30_000
INJECT_MAX_CHARS = 5_000

# Text-native форматы, читаемые как есть.
TEXT_EXT = {".md", ".txt", ".json", ".yaml", ".yml", ".csv", ".tsv"}

# Форматы, требующие извлечения (lazy-import) — пробуем и падаем graceful.
BINARY_EXT = {".pdf", ".docx"}

SUPPORTED_EXT = TEXT_EXT | BINARY_EXT


def _extract_pdf(raw: bytes) -> tuple[str, str]:
    """PDF → (text, note). Использует pypdf (pure-python, no system deps)."""
    try:
        import pypdf  # type: ignore
    except ImportError:
        return "", "pypdf not installed; pip install pypdf"
    try:
        import io as _io
        reader = pypdf.PdfReader(_io.BytesIO(raw))
        pages = []
        for i, p in enumerate(reader.pages):
            try:
                pages.append(p.extract_text() or "")
            except Exception as ex:
                pages.append(f"[page {i + 1}: extract failed: {ex}]")
        text = "\n\n".join(pages).strip()
        note = f"pdf: {len(reader.pages)} pages extracted"
        if not text:
            note += " (empty — possibly scanned/image PDF, OCR needed)"
        return text, note
    except Exception as ex:
        return "", f"pdf extract failed: {type(ex).__name__}: {ex}"


def _extract_docx(raw: bytes) -> tuple[str, str]:
    """DOCX → (text, note). Использует python-docx (paragraphs + tables)."""
    try:
        import docx  # type: ignore
    except ImportError:
        return "", "python-docx not installed; pip install python-docx"
    try:
        import io as _io
        doc = docx.Document(_io.BytesIO(raw))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        # tables — как простой tab-separated
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append("\t".join(cells))
        text = "\n".join(parts).strip()
        note = f"docx: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables"
        return text, note
    except Exception as ex:
        return "", f"docx extract failed: {type(ex).__name__}: {ex}"


def _decode_content(content: Any, ext: str) -> tuple[str, str]:
    """content → (text, extra_note). Handles str/bytes, JSON pretty, PDF/DOCX extract."""
    # str: text as-is (JSON pretty-print для .json)
    if isinstance(content, str):
        if ext == ".json":
            try:
                parsed = _json.loads(content)
                return _json.dumps(parsed, ensure_ascii=False, indent=2), ""
            except Exception:
                return content, ""
        return content, ""
    if isinstance(content, dict) or isinstance(content, list):
        try:
            return _json.dumps(content, ensure_ascii=False, indent=2), ""
        except Exception:
            return "", "content not json-serializable"
    if isinstance(content, bytes):
        if ext == ".pdf":
            return _extract_pdf(content)
        if ext == ".docx":
            return _extract_docx(content)
        # text-like binary → decode utf-8 (для .txt/.md upload'ов base64-декодированных)
        try:
            text = content.decode("utf-8", errors="replace")
            if ext == ".json":
                try:
                    parsed = _json.loads(text)
                    text = _json.dumps(parsed, ensure_ascii=False, indent=2)
                except Exception:
                    pass
            return text, ""
        except Exception:
            return "", "binary content not decodable as utf-8"
    return "", f"unsupported content type: {type(content).__name__}"


@dataclass
class NormalizedAttachment:
    filename: str
    ext: str
    text: str                # для инъекции (обрезано до INJECT_MAX_CHARS)
    full_text: str           # для audit (обрезано до ATTACH_MAX_CHARS)
    was_truncated: bool
    attach_to_persona: str = ""
    note: str = ""           # предупреждения если что-то не так


def _ext_of(filename: str) -> str:
    if not filename or "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def normalize(payload: dict[str, Any]) -> NormalizedAttachment | None:
    """Принимает payload attach_file intervention. Возвращает нормализованное
    представление или None если payload невалиден.

    Поддерживаемые input:
      - content str → text как есть (JSON pretty для .json)
      - content bytes → PDF/DOCX extractor или utf-8 decode
      - content dict/list → JSON serialize
      - content b64_content (str) + is_base64=True → decode → extract по ext
    """
    if not isinstance(payload, dict):
        return None
    filename = str(payload.get("filename") or "").strip()
    content = payload.get("content")
    b64_content = payload.get("b64_content")
    is_base64 = bool(payload.get("is_base64"))

    # base64-декодирование если запрошено (UI шлёт бинарные PDF/DOCX так)
    if b64_content and content is None:
        try:
            import base64 as _b64
            content = _b64.b64decode(b64_content, validate=False)
            is_base64 = True
        except Exception:
            return NormalizedAttachment(
                filename=filename or "unnamed", ext="", text="", full_text="",
                was_truncated=False,
                attach_to_persona=str(payload.get("attach_to_persona") or ""),
                note="b64 decode failed",
            )
    if content is None:
        return None
    attach_to_persona = str(payload.get("attach_to_persona") or "").strip()

    ext = _ext_of(filename)
    if ext and ext not in SUPPORTED_EXT:
        return NormalizedAttachment(
            filename=filename, ext=ext, text="", full_text="",
            was_truncated=False, attach_to_persona=attach_to_persona,
            note=f"unsupported ext {ext}; supported: {sorted(SUPPORTED_EXT)}",
        )

    text, decode_note = _decode_content(content, ext)
    if not text and decode_note:
        return NormalizedAttachment(
            filename=filename or "unnamed", ext=ext, text="", full_text="",
            was_truncated=False, attach_to_persona=attach_to_persona,
            note=decode_note,
        )

    was_truncated = False
    full_text = text
    if len(full_text) > ATTACH_MAX_CHARS:
        full_text = full_text[:ATTACH_MAX_CHARS] + "\n\n... [truncated]"
        was_truncated = True

    inject_text = full_text
    if len(inject_text) > INJECT_MAX_CHARS:
        inject_text = inject_text[:INJECT_MAX_CHARS] + "\n\n... [inject-clipped]"

    return NormalizedAttachment(
        filename=filename or "unnamed",
        ext=ext or ".txt",
        text=inject_text,
        full_text=full_text,
        was_truncated=was_truncated,
        attach_to_persona=attach_to_persona,
        note=decode_note,
    )


def format_attach_block(a: NormalizedAttachment) -> str:
    """Форматирует attach как блок для инъекции в system prompt."""
    header = f"## Прикреплённый пользователем материал\n"
    header += f"Файл: `{a.filename}`"
    if a.attach_to_persona:
        header += f" (для голоса `{a.attach_to_persona}`)"
    if a.was_truncated:
        header += " · **truncated**"
    return f"{header}\n\n```\n{a.text}\n```\n"
